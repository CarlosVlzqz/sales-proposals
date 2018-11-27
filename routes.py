from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from flask import render_template, request, Flask, redirect, url_for, session, send_file
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from num2words import num2words
from forms import Hardware, Software, Checkboxes, Servicios
from copy import deepcopy
import string
import os

app = Flask(__name__, static_url_path='')

port = int(os.getenv('PORT', 8000))

app.secret_key = "development-key"

#Search and replace function for each paragraph in the template
def Reemplazador(doc, palabras_clave, los_sustitutos):

    #Formatting of each user's input to title case or upper case if given case
    for i in range(len(los_sustitutos)):
        string_split = los_sustitutos[i].split()
        new_string = ""
        if i != 2:
            for j in range(len(string_split)):
                if string_split[j] == string_split[j].upper():
                    new_string += string_split[j] + " "
                else:
                    new_string += string_split[j].title() + " "
            los_sustitutos[i] = new_string[:(len(new_string)-1)]
        else:
            for j in range(len(string_split)):
                new_string += string_split[j]
                los_sustitutos[i] = new_string
 
    #Keyword iteration
    for i in range(0, len(palabras_clave), 1):
        #Paragraph Iteration
        for p in doc.paragraphs:
            #if keyword is found in paragraph
            if palabras_clave[i] in p.text:
                #Replacing the current keyword for the user's input
                p.text = p.text.replace(palabras_clave[i], los_sustitutos[i])
                #Applying "Normal" style to the paragraph
                p.style = doc.styles["Normal"]

    #Applying Title, Subtitle or Heading style when the variable is found
    for p in doc.paragraphs:
        if los_sustitutos[(len(los_sustitutos)-1)].upper() == p.text.upper():
            p.style = doc.styles["Title"]
        elif "subtitulo" in p.text:
            p.text = p.text.replace("(subtitulo)", "")
            p.style = doc.styles["Subtitle"]
        elif "Nuestra propuesta de valor" in p.text:
            p.style = doc.styles["Heading 2"]
        elif "Responsabilidades generales de IBM y " in p.text:
            p.style = doc.styles["Heading 1"]

#Search and replace in tables function
def Tablas(doc, los_sustitutos, palabras_clave, boole_boole):

    # Table index (needed for styling porpuses)
    num_table = 0
    for table in doc.tables:
        # Row index (needed for styling porpuses)
        num_row = 0
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for i in range(0, len(los_sustitutos), 1):
                        # If keyword is found, replace current keyword for the user's input
                        if palabras_clave[i] in p.text:
                            p.text = p.text.replace(palabras_clave[i], los_sustitutos[i])
                        # If it's indicated in the function call
                        if boole_boole == True:
                            # styling the first row of the first and third table
                            if num_row == 0 and (num_table == 2 or num_table == 0):
                                p.style = doc.styles["Subtitle"]
                        else:
                            # Otherwise apply the style only to the last table
                            if num_row == 0 and num_table == (len(doc.tables)-1):
                                p.style = doc.styles["Subtitle"]
            num_row += 1
        num_table += 1

# Function for deleting paragraphs
def delete_paragraph(paragraph):

    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# Function for copying tables in a docx file and inserting them in another file
def copy_table_after(table, paragraph):

    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)

# Function used to get the structure of a docx file
def iter_block_items(parent):

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

#List of spanish months names used to replace the month's number in the proposal
meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio",
         "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

#Main screen
@app.route("/")
def index():
    return render_template("index.html")

@app.route('/return-files-hw/')
def return_files_hw():
    try:
        return send_file('ContratoHW.docx', attachment_filename= f"{session.get('nombre_archivo', None)}", as_attachment=True)
    except Exception as e:
        return str(e)

@app.route('/return-files-sw/')
def return_files_sw():
    try:
        return send_file('ContratoSW.docx', attachment_filename= f"{session.get('nombre_archivo', None)}", as_attachment=True)
    except Exception as e:
        return str(e)

@app.route('/return-files-serv/')
def return_files_serv():
    try:
        return send_file('ContratoHW+S.docx', attachment_filename= f"{session.get('nombre', None)}", as_attachment=True)
    except Exception as e:
        return str(e)

#Hardware Form
@app.route("/hw", methods=['GET', 'POST'])
def hw():

    #Import of the hardware form
    form = Hardware()

    #POST method validation(Used to send data to a server)
    if request.method == 'POST':
        if form.validate() == False:
            # If it's not validated, reload the template and form
            return render_template("hardwareForm.html", form=form)
        # If it is valid:
        else:
            # Open the hardware template .docx file
            doc = Document("new_aa.docx")

            # Get and modify the first table
            table = doc.tables[0]
            table.cell(0,0).text = "Tipo"
            table.cell(0,1).text = "Descripción"
            table.cell(0,2).text = "Cantidad"

            # Create IBM style
            style = doc.styles["Normal"]
            font = style.font
            font.name = "IBM Plex Sans Text"
            font.size = Pt(10.5)

            # Title style
            estilo = doc.styles["Title"]
            fuente = estilo.font
            fuente.name = "IBM Plex Sans Text"
            fuente.size = Pt(25)
            fuente.color.rgb = RGBColor(0x1f, 0x4e, 0x79)

            # Subtitle style
            stylo = doc.styles["Subtitle"]
            fonto = stylo.font
            fonto.name = "IBM Plex Sans Light"
            fonto.size = Pt(12)
            fonto.color.rgb = RGBColor(0xff, 0xff, 0xff)
            fonto.bold = False

            # Heading style
            estilo = doc.styles["Heading 2"]
            fuente = estilo.font
            fuente.name = "IBM Plex Sans Light"
            fuente.bold = True

            # Get the form values
            nombre_vendedor = form.nombre_vendedor.data
            cargo_vendedor = form.cargo_vendedor.data
            mail_vendedor = form.mail_vendedor.data
            telefono_vendedor = form.telefono_vendedor.data
            contacto_cliente = form.contacto_cliente.data
            titulo_cliente = form.titulo_cliente.data
            razon_social = form.razon_social.data
            cliente_corto = form.cliente_corto.data
            estado_cliente = form.estado_cliente.data
            ciudad_cliente = form.ciudad_cliente.data
            colonia_cliente = form.colonia_cliente.data
            calle_cliente = form.calle_cliente.data
            numero_cliente = form.numero_cliente.data
            postal_cliente = form.postalCliente.data
            years_garantia = form.years_garantia.data + " (" + num2words(form.years_garantia.data, lang="es") + ")"
            precio_numero = str(form.precio_numero.data)
            fecha_contrato = str(form.fecha_contrato.data).split("-")
            dia_contrato = fecha_contrato[2]
            mes_contrato = meses[int(fecha_contrato[1])-1]
            year_contrato = fecha_contrato[0]
            fecha_corta = f"{dia_contrato}/{mes_contrato}/{year_contrato}"
            fecha_vigencia = str(form.fecha_vigencia.data).split("-")
            dia_vigencia = fecha_vigencia[2]
            mes_vigencia = meses[int(fecha_vigencia[1])-1]
            year_vigencia = fecha_vigencia[0]
            numero_propuesta = form.numero_propuesta.data
            titulo_contrato = form.titulo_contrato.data
            te_equis_te = form.config.data
            cargo_cliente = form.cargo_cliente.data

            if te_equis_te != None:
                #Read the .txt file
                lineas = te_equis_te.read()
                row = 1
                #Split it by lines
                array = str(lineas).split("\\n")
                #Clean the table
                for i in range(1, len(table.rows), 1):
                    table.cell(i ,0).text = ""
                #Clean special characters
                for i in range(0, len(array), 1):
                    string = array[i]
                    string = string.replace("\\r", "")
                    string = string.replace("\\n", "")
                    array[i] = string
                # Search for the article's key and insert it in the first table
                for i in range(0, len(array), 1):
                    indice = 0
                    arreglo = 0
                    string = array[i]
                    # Split the line if two spaces are found
                    temp = string.split("  ")
                    count = 0
                    # If the size of the array is higher than 2
                    if len(temp) > 2:
                        if len(temp[0]) == 8 and temp[0][4] == "-":
                            # Write the information of the file on each cell
                            while indice < 3:
                                if temp[arreglo] != "":
                                    try:
                                        table.cell(row, indice).text = temp[arreglo]
                                        arreglo += 1
                                        indice += 1
                                    # Add a row if the table seize is not enough
                                    except IndexError:
                                        table.add_row()
                                        table.cell(row, indice).text = temp[arreglo]
                                        arreglo += 1
                                        indice += 1
                                    # Delete unnecesary spaces
                                    try:
                                        if array[i+1][0:5] == "     " and count < 2:
                                            table.cell(row,1).text += array[i+1]
                                            table.cell(row,1).text = table.cell(row,1).text.replace("   ", "")
                                            count += 1
                                    except IndexError:
                                        pass
                                else:
                                    arreglo += 1
                            row += 1

            # Formatting of the price variable
            if len(str(precio_numero).split(".")[1]) == 1:
                precio_numero = str(precio_numero) + "0"

            # Getting the written value of the price, the cents, and formatting the price adding commas
            precio_letra = num2words(precio_numero.split(".")[0], lang="es").replace(" punto cero", "")
            centavos = precio_numero.split(".")[1]
            dolares = precio_numero.split(".")[0][::-1]
            precio_numero = ""
            for i in range(0, len(dolares), 3):
                precio_numero += dolares[i:i+3] + ","

            precio_numero = precio_numero[::-1][1:]

            years_garantia = years_garantia.replace(" punto cero", "")

            #Keywords to be searched in the file
            palabras_clave = ["nombreVendedor", "cargoVendedor", "mailVendedor",
                            "telefonoVendedor" , "contactoCliente", "razonSocial",
                            "clienteCorto", "cargoCliente", "estadoCliente", "ciudadCliente",
                            "coloniaCliente", "calleCliente", "numeroCliente",
                            "postalCliente", "yearsGarantia", "tituloCliente",
                            "precioNumero", "precioLetra", "centavos",
                            "diaContrato", "mesContrato", "yearContrato",
                            "diaVigencia", "mesVigencia", "yearVigencia",
                            "numeroPropuesta", "fechaCorta", "tituloContrato"]

            # User's inputs
            los_sustitutos = [nombre_vendedor, cargo_vendedor, mail_vendedor,
                            telefono_vendedor, contacto_cliente, razon_social,
                            cliente_corto, cargo_cliente, estado_cliente, ciudad_cliente,
                            colonia_cliente, calle_cliente, numero_cliente,
                            postal_cliente, years_garantia, titulo_cliente,
                            precio_numero, precio_letra, centavos,
                            dia_contrato, mes_contrato, year_contrato,
                            dia_vigencia, mes_vigencia, year_vigencia,
                            numero_propuesta.upper(), fecha_corta, titulo_contrato]

            # Replacing function call
            Reemplazador(doc, palabras_clave, los_sustitutos)

            # Deleting unused keywords and variables for the table replacing function
            palabras_clave = ["razonSocial", "calleCliente", "ciudadCliente", "postalCliente",
                              "estadoCliente", "numeroCliente", "coloniaCliente"]

            los_sustitutos = [razon_social.title(), calle_cliente.title(), ciudad_cliente.title(), postal_cliente,
                              estado_cliente.title(), numero_cliente, colonia_cliente.title()]

            #Table replacing function call
            Tablas(doc, los_sustitutos, palabras_clave, False)

            # Variable used to give a name to the file
            nombre_archivo = numero_propuesta.upper() + " - " + cliente_corto.upper()

            # Saving the changes to the file
            doc.save("ContratoHW.docx")

            session["nombre_archivo"] = f"{cliente_corto.title()} - {numero_propuesta.upper()}.docx"

            return redirect(url_for('return_files_hw'))

    elif request.method == "GET":
        return render_template("hardwareForm.html", form=form)

#Software Form
@app.route("/sw", methods=["GET", "POST"])
def sw():

    # Import software form
    form = Software()

    # Post method validation
    if request.method == "POST":
        if form.validate() == False:
            return render_template("softwareForm.html", form=form)

        else:
            # Open software template file
            doc = Document("new_bb.docx")

            #IBM style
            style = doc.styles["Normal"]
            font = style.font
            font.name = "IBM Plex Sans Text"
            font.size = Pt(10.5)

            # Tile Style
            style = doc.styles["Title"]
            font = style.font
            font.name = "IBM Plex Sans Text"
            font.size = Pt(28.5)
            font.color.rgb = RGBColor(0xff, 0xff, 0xff)

            # Subtitle style
            stylo = doc.styles["Subtitle"]
            fonto = stylo.font
            fonto.name = "IBM Plex Sans Light"
            fonto.size = Pt(10.5)
            fonto.color.rgb = RGBColor(0xff, 0xff, 0xff)
            fonto.bold = True

            # Heading style
            estilo = doc.styles["Heading 2"]
            fuente = estilo.font
            fuente.name = "IBM Plex Sans Light"
            fuente.bold = True

            # Get the form values
            nombre_vendedor = form.nombre_vendedor.data
            cargo_vendedor = form.cargo_vendedor.data
            mail_vendedor = form.mail_vendedor.data
            telefono_vendedor = form.telefono_vendedor.data
            contacto_cliente = form.contacto_cliente.data
            titulo_cliente = form.titulo_cliente.data
            razon_social = form.razon_social.data
            cliente_corto = form.cliente_corto.data
            estado_cliente = form.estado_cliente.data
            ciudad_cliente = form.ciudad_cliente.data
            colonia_cliente = form.colonia_cliente.data
            calle_cliente = form.calle_cliente.data
            numero_cliente = form.numero_cliente.data
            postal_cliente = form.postalCliente.data
            precio_numero = form.precio_numero.data
            fecha_contrato = str(form.fecha_contrato.data).split("-")
            dia_contrato = fecha_contrato[2]
            mes_contrato = meses[int(fecha_contrato[1])-1]
            year_contrato = fecha_contrato[0]
            fecha_vigencia = str(form.fecha_vigencia.data).split("-")
            dia_vigencia = fecha_vigencia[2]
            mes_vigencia = meses[int(fecha_vigencia[1])-1]
            year_vigencia = fecha_vigencia[0]
            numero_propuesta = form.numero_propuesta.data.upper()
            titulo_contrato = form.titulo_contrato.data
            to_addr = form.to_addr.data
            pe_de_efe = form.config.data
            cargo_cliente = form.cargo_cliente.data

            table = doc.tables[0]

            # Validating a PDF file has been loaded
            if pe_de_efe != None:
                pe_de_efe = Document(pe_de_efe)
                # Validating a PDF file has been loaded
                for table in pe_de_efe.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if len(cell.paragraphs) > 1:
                                parrafos = list(cell.paragraphs)
                                break

                row = 1

                try:

                    for i in range(len(parrafos)):
                        if parrafos[i].text[:3] == f"00{row}":
                            tabla_sw = doc.tables[0]
                            try:
                                tabla_sw.cell(row, 1).text = parrafos[i-1].text
                                string = parrafos[i].text
                                split = string.split()
                                tabla_sw.cell(row, 0).text = split[1]
                                tabla_sw.cell(row, 2).text = split[2]
                                string = parrafos[i+1].text
                                split = string.split(" - ")
                                tabla_sw.cell(row, 3).text = split[0]
                                tabla_sw.cell(row, 4).text = split[1]
                                row += 1
                            except IndexError:
                                tabla_sw.add_row()
                                tabla_sw.cell(row, 1).text = parrafos[i-1].text
                                string = parrafos[i].text
                                split = string.split()
                                tabla_sw.cell(row, 0).text = split[1]
                                tabla_sw.cell(row, 2).text = split[2]
                                string = parrafos[i+1].text
                                split = string.split(" - ")
                                tabla_sw.cell(row, 3).text = split[0]
                                tabla_sw.cell(row, 4).text = split[1]
                                row += 1

                except UnboundLocalError:

                    row = 1

                    parrafos = list(pe_de_efe.paragraphs)

                    for i in range(len(parrafos)):
                        if f"00{row}" in parrafos[i].text:
                            print(parrafos[i].text)
                            tabla_sw = doc.tables[0]
                            try:
                                tabla_sw.cell(row, 1).text = parrafos[i-1].text
                                string = parrafos[i].text
                                split = string.split()
                                tabla_sw.cell(row, 0).text = split[1]
                                tabla_sw.cell(row, 2).text = split[2]
                                string = parrafos[i+1].text
                                split = string.split(" - ")
                                tabla_sw.cell(row, 3).text = split[0]
                                tabla_sw.cell(row, 4).text = split[1]
                                row += 1
                            except IndexError:
                                tabla_sw.add_row()
                                tabla_sw.cell(row, 1).text = parrafos[i-1].text
                                string = parrafos[i].text
                                split = string.split()
                                tabla_sw.cell(row, 0).text = split[1]
                                tabla_sw.cell(row, 2).text = split[2]
                                string = parrafos[i+1].text
                                split = string.split(" - ")
                                tabla_sw.cell(row, 3).text = split[0]
                                tabla_sw.cell(row, 4).text = split[1]
                                row += 1

            else: 
                pass

            # Formatting the price input
            if len(str(precio_numero).split(".")[1]) == 1:
                precio_numero = str(precio_numero) + "0"


            # Getting the written value of the price, the cents, and formatting the price adding commas
            precio_letra = num2words(precio_numero.split(".")[0], lang="es").replace(" punto cero", "")
            centavos = precio_numero.split(".")[1]
            dolares = precio_numero.split(".")[0][::-1]
            precio_numero = ""
            for i in range(0, len(dolares), 3):
                precio_numero += dolares[i:i+3] + ","

            precio_numero = precio_numero[::-1][1:]

            if "Sa De Cv" in razon_social or "S A De C V" in razon_social:
                razon_social = razon_social.replace("Sa De Cv", "S.A. de C.V.")
                razon_social = razon_social.replace("s a de c v", "S.A. de C.V.")

            # Keywords to be searched in the template
            palabras_clave = ["nombreVendedor", "cargoVendedor", "mailVendedor",
                            "telefonoVendedor", "cargoCliente", "contactoCliente", "razonSocial",
                            "clienteCorto", "estadoCliente", "ciudadCliente",
                            "coloniaCliente", "calleCliente", "numeroCliente",
                            "postalCliente", "tituloCliente",
                            "precioNumero", "precioLetra", "centavos",
                            "diaContrato", "mesContrato", "yearContrato",
                            "diaVigencia", "mesVigencia", "yearVigencia",
                            "numeroPropuesta", "tituloContrato"]

            # Variables of the form
            los_sustitutos = [nombre_vendedor, cargo_vendedor, mail_vendedor,
                            telefono_vendedor, cargo_cliente, contacto_cliente, razon_social,
                            cliente_corto, estado_cliente, ciudad_cliente,
                            colonia_cliente, calle_cliente, numero_cliente,
                            postal_cliente, titulo_cliente,
                            precio_numero, precio_letra, centavos,
                            dia_contrato, mes_contrato, year_contrato,
                            dia_vigencia, mes_vigencia, year_vigencia,
                            numero_propuesta, titulo_contrato]

            # Replacing function call
            Reemplazador(doc, palabras_clave, los_sustitutos)

            # Cleaning unused variables for the table replacing function
            palabras_clave = ["razonSocial", "calleCliente", "ciudadCliente", "postalCliente",
                              "estadoCliente", "numeroCliente", "coloniaCliente"]

            los_sustitutos = [razon_social.title(), calle_cliente.title(), ciudad_cliente.title(), postal_cliente,
                              estado_cliente.title(), numero_cliente, colonia_cliente.title()]

            # Table replacing function
            Tablas(doc, los_sustitutos, palabras_clave, True)

            # Variable used to name the file
            nombre_archivo = numero_propuesta.upper()+ " - " + cliente_corto.upper()

            # #Guardar los cambios al documento
            doc.save("ContratoSW.docx")

            session["nombre_archivo"] = f"{cliente_corto.title()} - {numero_propuesta.upper()}.docx"

            return redirect(url_for('return_files_sw'))

    elif request.method == "GET":
        return render_template("softwareForm.html", form=form)

# HW + Services form
@app.route("/serv", methods=["GET", "POST"])
def serv():

    # Import form
    form = Servicios()

    # POST Validation
    if request.method == "POST":
        if form.validate() == False:
            return render_template("servicios.html", form=form)

        else:
            # Open file
            doc = Document("new_cc.docx")

            # Get and edit the first table
            table = doc.tables[0]
            table.cell(0,0).text = "Tipo"
            table.cell(0,1).text = "Descripción"
            table.cell(0,2).text = "Cantidad"

            # IBM Style
            style = doc.styles["Normal"]
            font = style.font
            font.name = "IBM Plex Sans Text"
            font.size = Pt(10.5)

            # Title style
            estilo = doc.styles["Title"]
            fuente = estilo.font
            fuente.name = "IBM Plex Sans Text"
            fuente.size = Pt(36)
            fuente.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # Subtitle style
            stylo = doc.styles["Subtitle"]
            fonto = stylo.font
            fonto.name = "IBM Plex Sans Light"
            fonto.size = Pt(12)
            fonto.color.rgb = RGBColor(0xff, 0xff, 0xff)
            fonto.bold = False

            # Geading style
            estilo = doc.styles["Heading 2"]
            fuente = estilo.font
            fuente.name = "IBM Plex Sans Light"
            fuente.bold = True

            # Get form values
            nombre_vendedor = form.nombre_vendedor.data
            cargo_vendedor = form.cargo_vendedor.data
            mail_vendedor = form.mail_vendedor.data
            telefono_vendedor = form.telefono_vendedor.data
            contacto_cliente = form.contacto_cliente.data
            titulo_cliente = form.titulo_cliente.data
            razon_social = form.razon_social.data
            cliente_corto = form.cliente_corto.data
            estado_cliente = form.estado_cliente.data
            ciudad_cliente = form.ciudad_cliente.data
            colonia_cliente = form.colonia_cliente.data
            calle_cliente = form.calle_cliente.data
            numero_cliente = form.numero_cliente.data
            postal_cliente = form.postalCliente.data
            years_garantia = form.years_garantia.data + " (" + num2words(form.years_garantia.data, lang="es") + ")"
            precio_numero = form.precio_numero.data
            fecha_contrato = str(form.fecha_contrato.data).split("-")
            dia_contrato = fecha_contrato[2]
            mes_contrato = meses[int(fecha_contrato[1])-1]
            year_contrato = fecha_contrato[0]
            fecha_vigencia = str(form.fecha_vigencia.data).split("-")
            dia_vigencia = fecha_vigencia[2]
            mes_vigencia = meses[int(fecha_vigencia[1])-1]
            year_vigencia = fecha_vigencia[0]
            numero_propuesta = form.numero_propuesta.data
            titulo_contrato = form.titulo_contrato.data
            # to_addr = form.to_addr.data
            te_equis_te = form.config.data
            cargo_cliente = form.cargo_cliente.data

            # Validate config file's been uploaded
            if te_equis_te != None:
                #Read the .txt file
                lineas = te_equis_te.read()
                row = 1
                #Split it by lines
                array = str(lineas).split("\\n")
                #Clean the table
                for i in range(1, len(table.rows), 1):
                    table.cell(i ,0).text = ""
                #Clean special characters
                for i in range(0, len(array), 1):
                    string = array[i]
                    string = string.replace("\\r", "")
                    string = string.replace("\\n", "")
                    array[i] = string
                # Search for the article's key and insert it in the first table
                for i in range(0, len(array), 1):
                    indice = 0
                    arreglo = 0
                    string = array[i]
                    # Split the line if two spaces are found
                    temp = string.split("  ")
                    count = 0
                    # If the size of the array is higher than 2
                    if len(temp) > 2:
                        if len(temp[0]) == 8 and temp[0][4] == "-":
                            # Write the information of the file on each cell
                            while indice < 3:
                                if temp[arreglo] != "":
                                    try:
                                        table.cell(row, indice).text = temp[arreglo]
                                        arreglo += 1
                                        indice += 1
                                    # Add a row if the table seize is not enough
                                    except IndexError:
                                        table.add_row()
                                        table.cell(row, indice).text = temp[arreglo]
                                        arreglo += 1
                                        indice += 1
                                    # Delete unnecesary spaces
                                    try:
                                        if array[i+1][0:5] == "     " and count < 2:
                                            table.cell(row,1).text += array[i+1]
                                            table.cell(row,1).text = table.cell(row,1).text.replace("   ", "")
                                            count += 1
                                    except IndexError:
                                        pass
                                else:
                                    arreglo += 1
                            row += 1

            # format price
            if len(str(precio_numero).split(".")[1]) == 1:
                precio_numero = str(precio_numero) + "0"

            precio_numero = str(precio_numero)

            # Getting the written value of the price, the cents, and formatting the price adding commas
            precio_letra = num2words(str(precio_numero).split(".")[0], lang="es").replace(" punto cero", "")
            centavos = precio_numero.split(".")[1]
            dolares = precio_numero.split(".")[0][::-1]
            precio_numero = ""
            for i in range(0, len(dolares), 3):
                precio_numero += dolares[i:i+3] + ","

            precio_numero = precio_numero[::-1][1:]

            years_garantia = years_garantia.replace(" punto cero", "")

            #Pablabras clave a buscar en el documento
            palabras_clave = ["nombreVendedor", "cargoVendedor", "mailVendedor",
                            "telefonoVendedor", "cargoCliente", "contactoCliente", "razonSocial",
                            "clienteCorto", "estadoCliente", "ciudadCliente",
                            "coloniaCliente", "calleCliente", "numeroCliente",
                            "postalCliente", "yearsGarantia", "tituloCliente",
                            "precioNumero", "precioLetra", "centavos",
                            "diaContrato", "mesContrato", "yearContrato",
                            "diaVigencia", "mesVigencia", "yearVigencia",
                            "numeroPropuesta", "tituloContrato"]

            los_sustitutos = [nombre_vendedor, cargo_vendedor, mail_vendedor,
                            telefono_vendedor, cargo_cliente, contacto_cliente, razon_social,
                            cliente_corto, estado_cliente, ciudad_cliente,
                            colonia_cliente, calle_cliente, numero_cliente,
                            postal_cliente, years_garantia, titulo_cliente,
                            precio_numero, precio_letra, centavos,
                            dia_contrato, mes_contrato, year_contrato,
                            dia_vigencia, mes_vigencia, year_vigencia,
                            numero_propuesta.upper(), titulo_contrato]

            # Replacing function call
            Reemplazador(doc, palabras_clave, los_sustitutos)

            # Cleaning unused variables
            palabras_clave = ["razonSocial", "calleCliente", "ciudadCliente", "postalCliente",
                              "estadoCliente", "numeroCliente", "coloniaCliente"]

            # Naming the proposal
            nombre_archivo = numero_propuesta.upper() + " - " + cliente_corto.upper()

            # Table variables
            session["table_variables"] = f"{razon_social}-{calle_cliente}-{ciudad_cliente}-{postal_cliente}-{estado_cliente}-{numero_cliente}-{colonia_cliente}"

            # If OTC was selected
            if form.OTC.data == "OTC":
                # Delete mothly payment paragraph
                doc.paragraphs[158].text = doc.paragraphs[158].text.replace(" y Servicios", "")
                delete_paragraph(doc.paragraphs[159])
            else:
                pass

            doc.save("ContratoHW+S.docx") 

            session["nombre_archivo"] = f"{cliente_corto.title()} - {numero_propuesta.upper()}.docx"   

            return redirect(url_for('check'))

    elif request.method == "GET":
        return render_template("servicios.html", form=form)

#Segunda pagina de HW + Servicios (Seleccion de servicios)
@app.route("/check", methods=["GET", "POST"])
def check():

    # import the form
    form = Checkboxes()

    # Validate POST method
    if request.method == "POST":
        if form.validate() == False:
            return render_template("check.html", form=form)

        else:
            # Open the file
            doc = Document("ContratoHW+S.docx")

            session["nombre"] = session.get('nombre_archivo', None)
            table_variables = session.get('table_variables', None)

            # Create style
            style = doc.styles["Normal"]
            font = style.font
            font.name = "IBM Plex Sans Text"
            font.size = Pt(10.5)
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            # font.bold = True

            stylo = doc.styles["Heading 2"]
            fonto = stylo.font
            fonto.name = "IBM Plex Sans Text"
            fonto.size = Pt(11)
            fonto.bold = True

            palabras_clave = ["razonSocial", "calleCliente", "ciudadCliente", "postalCliente",
                              "estadoCliente", "numeroCliente", "coloniaCliente"]

            los_sustitutos = table_variables.split("-")

            headings = ["Servicios de LAB Services", "Servicios Profesionales de IBM", "Servicios de Mantenimiento"]

            # Description of each service
            leyendas = ["La presente Propuesta incluye los siguientes servicios especializados de tecnología para los equipos descritos en esta Propuesta:",
                        "Dentro del alcance de esta Propuesta se consideran los siguientes servicios:",
                        "La presente Propuesta incluye los siguientes servicios de mantenimiento para los equipos descritos en esta Propuesta:"]

            # Heading of anex
            servicios = ["\"Alcance de servicios de LAB services\"", "\"Servicios profesionales\"",  "\"Servicios de mantenimiento\""]

            letras = ["A", "B", "C"]

            apartado_firmas = ["Apartado de firmas", 
                               "Mediante la firma de la presente Propuesta, las Partes, cuyas denominaciones y apoderados aparecen al calce de este instrumento (en lo sucesivo “Las Partes”), declaran haber leído y entendido los términos y condiciones del Acuerdo de Relación con Cliente publicado en la siguiente página web: https://www-05.ibm.com/support/operations/mx/es/documents.html (en adelante, el “Contrato”) aceptando que el mismo se aplica por incorporación a, y forma parte integrante, de esta Propuesta y de sus anexos, como si estuviese transcrito a la letra en este documento.",
                               "Asimismo, los apoderados de las Partes declaran que tienen facultades suficientes para la suscripción de esta Propuesta y sus anexos, y que dichas facultades no les han sido limitadas ni de cualquier manera modificadas o revocadas a la fecha de firma de la presente.", 
                               "FIN DEL DOCUMENTO",]

            i=0

            for p in doc.paragraphs:
                if "Espacio intencionalmente dejado en blanco" in p.text:
                    break
                i += 1

            checks = [form.serv1.data, form.serv2.data, form.serv3.data]

            index = [i for i, j in enumerate(checks) if j == True]

            for p in range(checks.count(True)):
                doc.paragraphs[i].insert_paragraph_before(headings[index[p]])
                doc.paragraphs[i].style = doc.styles["Heading 2"]
                i += 1
                doc.paragraphs[i].insert_paragraph_before(leyendas[index[p]])
                i += 1
                doc.paragraphs[i].insert_paragraph_before("Ejemplos: ")
                doc.paragraphs[i].style = doc.styles["List Paragraph"]
                i += 1
                doc.paragraphs[i].insert_paragraph_before(f"En el Anexo {letras[p]} de este documento se describe el alcance y cobertura de este servicio.")
                i += 1
                
            doc.paragraphs[i].insert_paragraph_before()
            doc.paragraphs[i].insert_paragraph_before()
            doc.paragraphs[i].insert_paragraph_before()
            doc.paragraphs[i].insert_paragraph_before()

            i = -1

            for p in doc.paragraphs:
                if "Apartado de firmas" in p.text:
                    break
                i += 1

            for p in range(checks.count(True)):
                paragraph = doc.add_paragraph(f"Anexo {letras[p]} {servicios[index[p]]}")
                paragraph.style = doc.styles["Heading 5"]
                i += 1
                doc.add_paragraph()
                i += 1

            if form.serv3.data == True:

                doc.add_paragraph(f"Documento de Transacción Número: {form.doc_transaccion.data}")
                doc.add_paragraph()

                doc2 = Document(form.config.data)

                estructura = []

                for block in iter_block_items(doc2):
                    if block.__class__.__name__ == "Paragraph":
                        estructura.append("P")
                    else:
                        estructura.append("T")

                tablas = 2
                count = 0

                for item in range(169, len(estructura), 1):
                    if estructura[item] == "P":
                        try:
                            doc.add_paragraph(doc2.paragraphs[item].text)
                            print(doc2.paragraphs[item].text)
                        except IndexError:
                            pass
                    else:
                        try:
                            copy_table_after(doc2.tables[tablas], doc.paragraphs[len(doc.paragraphs)-1])
                        except IndexError:
                            copy_table_after(doc2.tables[tablas], doc.paragraphs[len(doc.paragraphs)-1])
                        tablas += 1
                        count += 1

            doc3 = Document("new_aa.docx")

            paragraph = doc.add_paragraph(apartado_firmas[0])
            paragraph.style = doc.styles["Heading 1"]
            paragraph = doc.add_paragraph(apartado_firmas[1])
            paragraph.style = doc.styles["Normal"]
            paragraph = doc.add_paragraph(apartado_firmas[2])
            paragraph.style = doc.styles["Normal"]
            paragraph = doc.add_paragraph(apartado_firmas[3])
            paragraph.style = doc.styles["Normal"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            copy_table_after(doc3.tables[2], doc.paragraphs[len(doc.paragraphs)-2])

            # Table replacing function
            Tablas(doc, los_sustitutos, palabras_clave, True)
            
            doc.save("ContratoHW+S.docx")

            return redirect(url_for('return_files_serv'))

    elif request.method == "GET":
        return render_template("check.html", form=form)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=port, debug=True)