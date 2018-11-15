from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from flask import render_template, request, Flask, redirect, url_for, session, send_file
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from num2words import num2words
from forms import Hardware, Software, Checkboxes, Servicios
from copy import deepcopy
import string
import os

def delete_paragraph(paragraph):

    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_table_after(table, paragraph):

    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)

def iter_block_items(parent):

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

doc = Document("new_cc.docx")

headings = ["Servicios de LAB Services", "Servicios Profesionales de IBM", "Servicios de Mantenimiento"]

# Description of each service
leyendas = ["La presente Propuesta incluye los siguientes servicios especializados de tecnología para los equipos descritos en esta Propuesta:",
            "Dentro del alcance de esta Propuesta se consideran los siguientes servicios:",
            "La presente Propuesta incluye los siguientes servicios de mantenimiento para los equipos descritos en esta Propuesta:"]

# Heading of anex
servicios = ["\"Alcance de servicios de LAB services\"", "\"Servicios profesionales\"",  "\"Servicios de mantenimiento\""]

letras = ["A", "B", "C"]

i=0

for p in doc.paragraphs:
    if "Espacio intencionalmente dejado en blanco" in p.text:
        break
    i += 1

checks = [True, True, True]

index = [i for i, j in enumerate(checks) if j == True]

for p in range(checks.count(True)):
    doc.paragraphs[i].insert_paragraph_before(headings[index[p]])
    doc.paragraphs[i].style = doc.styles["Heading 2"]
    i += 1
    doc.paragraphs[i].insert_paragraph_before(leyendas[index[p]])
    i += 1
    doc.paragraphs[i].insert_paragraph_before("Ejemplos: ")
    doc.paragraphs[i].style = doc.styles["List Paragraph"]
    i += 1
    doc.paragraphs[i].insert_paragraph_before(f"En el Anexo {letras[p]} de este documento se describe el alcance y cobertura de este servicio.")
    i += 1
	
doc.paragraphs[i].insert_paragraph_before()
doc.paragraphs[i].insert_paragraph_before()
doc.paragraphs[i].insert_paragraph_before()
doc.paragraphs[i].insert_paragraph_before()

i = 0

for p in doc.paragraphs:
    if "Apartado de firmas" in p.text:
        break
    i += 1

for p in range(checks.count(True)):
    doc.paragraphs[i].insert_paragraph_before(f"Anexo {letras[p]} {servicios[index[p]]}")
    doc.paragraphs[i].style = doc.styles["Heading 5"]
    i += 1
    doc.paragraphs[i].insert_paragraph_before()
    i += 1

doc2 = Document("chis.docx")

estructura = []

for block in iter_block_items(doc2):
    if block.__class__.__name__ == "Paragraph":
        estructura.append("P")
    else:
        estructura.append("T")

new = Document()
tablas = 0

doc.tables[len(doc.tables)-1]._element.getparent().remove(doc.tables[len(doc.tables)-1]._element)
delete_paragraph(doc.paragraphs[len(doc.paragraphs)-1])

for item in range(len(estructura)):
    try:
        if estructura[item] == "P":
            new.add_paragraph(doc2.paragraphs[item].text)
        else:
            copy_table_after(doc2.tables[tablas], new.paragraphs[len(new.paragraphs)-5])
            tablas += 1
    except IndexError:
        pass

new.add_paragraph("Firmas, firmas")
new.add_paragraph("Bla bla bla bla firmas")
new.add_paragraph("Bla bla bla bla firmas y mas firmas")
copy_table_after(doc.tables[len(doc.tables)-1], new.paragraphs[len(new.paragraphs)-1])
new.add_paragraph("Fin")

new.save("new.docx")

doc.tables[len(doc.tables)-1]._element.getparent().remove(doc.tables[len(doc.tables)-1]._element)

doc.save("result.docx")